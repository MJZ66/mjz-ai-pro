"""多格式文件读取、图片编码与文本截断。"""
import base64
import io
from dataclasses import dataclass
from pathlib import Path
from typing import Any, List, Optional, Tuple

from pypdf import PdfReader

from utils.logger import get_logger, user_friendly_error

logger = get_logger(__name__)

DEFAULT_MAX_FILE_CHARS = 12000
TRUNCATE_NOTICE = "已截断，仅使用前 {max_chars} 字符参与对话。"

TEXT_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".xlsx", ".xls"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | IMAGE_EXTENSIONS

# Streamlit file_uploader type 列表
UPLOAD_TYPES = [
    "txt", "md", "pdf", "docx", "xlsx", "xls", "png", "jpg", "jpeg",
]

IMAGE_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
}


class FileLoadError(Exception):
    """文件加载失败。"""


@dataclass
class LoadedFile:
    filename: str
    extension: str
    kind: str  # text | image
    text: str = ""
    image_base64: str = ""
    image_mime: str = "image/png"
    size_bytes: int = 0

    @property
    def is_image(self) -> bool:
        return self.kind == "image"

    @property
    def is_text(self) -> bool:
        return self.kind == "text"


def truncate_text(
    text: str,
    max_chars: int = DEFAULT_MAX_FILE_CHARS,
) -> Tuple[str, bool]:
    if not text:
        return "", False
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


def prepare_file_content(
    text: str,
    max_chars: int = DEFAULT_MAX_FILE_CHARS,
) -> Tuple[str, bool, int]:
    original_len = len(text or "")
    truncated, was_truncated = truncate_text(text or "", max_chars)
    return truncated, was_truncated, original_len


def truncate_notice(max_chars: int = DEFAULT_MAX_FILE_CHARS) -> str:
    return TRUNCATE_NOTICE.format(max_chars=max_chars)


def load_uploaded_file(uploaded_file: Any) -> LoadedFile:
    """从 Streamlit UploadedFile 解析为 LoadedFile。"""
    if uploaded_file is None:
        raise FileLoadError("未选择文件")

    name = getattr(uploaded_file, "name", "unknown")
    raw = uploaded_file.read()
    if isinstance(raw, str):
        raw = raw.encode("utf-8")
    mime = getattr(uploaded_file, "type", None)
    return load_from_bytes(raw, name, mime_type=mime)


def load_from_bytes(
    data: bytes,
    filename: str,
    *,
    mime_type: Optional[str] = None,
) -> LoadedFile:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise FileLoadError(
            f"不支持的文件类型：{suffix}。"
            f"支持：{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    try:
        if suffix in IMAGE_EXTENSIONS:
            return _load_image(data, filename, suffix)
        text = _load_text_bytes(data, filename, suffix, mime_type)
        return LoadedFile(
            filename=filename,
            extension=suffix,
            kind="text",
            text=text,
            size_bytes=len(data),
        )
    except FileLoadError:
        raise
    except Exception as exc:
        logger.exception("文件解析失败: %s", filename)
        raise FileLoadError(user_friendly_error(exc)) from exc


def _load_text_bytes(
    data: bytes,
    filename: str,
    suffix: str,
    mime_type: Optional[str],
) -> str:
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="replace")
    if suffix == ".pdf" or mime_type == "application/pdf":
        return _load_pdf(data)
    if suffix == ".docx":
        return _load_docx(data)
    if suffix in {".xlsx", ".xls"}:
        return _load_excel(data, suffix)
    if mime_type == "text/plain":
        return data.decode("utf-8", errors="replace")
    raise FileLoadError(f"无法解析文本文件：{filename}")


def _load_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "".join(parts)


def _load_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise FileLoadError("请安装 python-docx：pip install python-docx") from exc

    doc = Document(io.BytesIO(data))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_excel(data: bytes, suffix: str) -> str:
    try:
        import openpyxl
    except ImportError as exc:
        raise FileLoadError("请安装 openpyxl：pip install openpyxl") from exc

    if suffix == ".xls":
        try:
            import xlrd
        except ImportError as exc:
            raise FileLoadError(
                "读取 .xls 需要安装 xlrd：pip install xlrd"
            ) from exc
        book = xlrd.open_workbook(file_contents=data)
        parts = []
        for sheet in book.sheets():
            parts.append(f"## Sheet: {sheet.name}")
            for rx in range(sheet.nrows):
                row = [
                    str(sheet.cell_value(rx, cx)).strip()
                    for cx in range(sheet.ncols)
                ]
                if any(row):
                    parts.append(" | ".join(row))
        return "\n".join(parts)

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def _load_image(data: bytes, filename: str, suffix: str) -> LoadedFile:
    meta = ""
    try:
        from PIL import Image

        img = Image.open(io.BytesIO(data))
        meta = f"尺寸 {img.size[0]}×{img.size[1]}，格式 {img.format or suffix}"
    except Exception:
        meta = f"图片文件 {suffix}"

    b64 = base64.standard_b64encode(data).decode("ascii")
    return LoadedFile(
        filename=filename,
        extension=suffix,
        kind="image",
        text=f"[图片附件] {filename}（{meta}）。请结合用户问题分析图片内容。",
        image_base64=b64,
        image_mime=IMAGE_MIME.get(suffix, "image/png"),
        size_bytes=len(data),
    )


def build_file_user_message(
    loaded: LoadedFile,
    user_note: str = "",
    *,
    max_chars: int = DEFAULT_MAX_FILE_CHARS,
) -> dict:
    """构建可送入 LLM 的 user 消息（支持图片多模态）。"""
    note = (user_note or "").strip()
    prefix = f"{note}\n\n" if note else ""

    if loaded.is_image:
        text_part = (
            f"{prefix}请分析附件图片（{loaded.filename}）。"
            "描述画面关键信息并回答用户问题。"
        )
        return {
            "role": "user",
            "content": [
                {"type": "text", "text": text_part},
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{loaded.image_mime};base64,{loaded.image_base64}",
                    },
                },
            ],
        }

    body, truncated, _ = prepare_file_content(loaded.text, max_chars)
    trunc_hint = f"\n\n（{truncate_notice(max_chars)}）" if truncated else ""
    content = (
        f"{prefix}请分析以下文件：\n"
        f"文件名：{loaded.filename}\n"
        f"文件类型：{loaded.extension}\n\n"
        f"文件内容：\n{body}{trunc_hint}"
    )
    return {"role": "user", "content": content}


def format_message_display(content: Any) -> str:
    """将消息内容格式化为页面可展示的字符串。"""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        texts = []
        has_image = False
        for part in content:
            if not isinstance(part, dict):
                continue
            if part.get("type") == "text":
                texts.append(part.get("text", ""))
            elif part.get("type") == "image_url":
                has_image = True
        display = "\n".join(t for t in texts if t)
        if has_image:
            display += "\n\n🖼️ [已附带图片，使用视觉模型时可分析]"
        return display.strip() or "[多媒体消息]"
    return str(content)


# 兼容旧接口
def read_uploaded_file(uploaded_file: Any) -> str:
    loaded = load_uploaded_file(uploaded_file)
    return loaded.text
